from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

from docx import Document as DocxDocument  # python-docx

from app.documents.models import Document


def extract_docx_text(path: str | Path) -> str:
    """
    Извлекает текст из DOCX-документа (все параграфы подряд).
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"DOCX not found: {p}")

    doc = DocxDocument(str(p))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n\n".join(paragraphs).strip()


def load_docx_document(path: str | Path, doc_id: str | None = None, title: str | None = None) -> Tuple[
    Document, List[str]]:
    """
    Загружает DOCX как Document + список "страниц".
    Пока считаем, что весь текст — одна логическая страница (index 0).

    Возвращает:
      - Document
      - список из одной строки [full_text]
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"DOCX not found: {p}")

    full_text = extract_docx_text(p)

    if doc_id is None:
        doc_id = p.stem

    if title is None:
        title = p.stem

    doc = Document(
        id=doc_id,
        path=p,
        title=title,
        pages=1,
    )

    return doc, [full_text]
